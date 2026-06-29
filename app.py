"""
AI Marketing Proposal Generator
================================
Streamlit frontend — smart form with multiple output types,
language selection, and dynamic fields per document type.
"""

from __future__ import annotations

import io
import logging
import os
import sys
import traceback
from pathlib import Path

import streamlit as st
from docx import Document as DocxDocument

sys.path.insert(0, str(Path(__file__).parent))

import config
from services.proposal_service import ProposalService

logging.basicConfig(
    level=getattr(logging, config.LOG_LEVEL, logging.INFO),
    format="%(asctime)s | %(levelname)-8s | %(name)s | %(message)s",
)
logger = logging.getLogger(__name__)

st.set_page_config(
    page_title="AI Marketing Proposal Generator",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
.main-header {
    background: linear-gradient(135deg, #1e3a5f 0%, #2d6a9f 100%);
    color: white; padding: 1.5rem 2rem; border-radius: 10px; margin-bottom: 1.5rem;
}
.score-high { color: #16a34a; font-weight: bold; }
.score-mid  { color: #d97706; font-weight: bold; }
.score-low  { color: #dc2626; font-weight: bold; }
.doc-type-badge {
    background: #e8f0fe; color: #1a56db; padding: 2px 10px;
    border-radius: 20px; font-size: 0.85rem; font-weight: 600;
}
</style>
""", unsafe_allow_html=True)

@st.cache_resource(show_spinner="Initialising AI services…")
def get_service() -> ProposalService:
    return ProposalService()

service = get_service()

if "chat_history" not in st.session_state:
    st.session_state.chat_history: list[dict] = []
if "last_retrieved_chunks" not in st.session_state:
    st.session_state.last_retrieved_chunks: list[dict] = []
if "last_context_sent" not in st.session_state:
    st.session_state.last_context_sent: str = ""

# ══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════════════════════════════

with st.sidebar:
    st.markdown("## 📁 Document Management")
    st.markdown("### Upload Documents")
    doc_type = st.selectbox("Document Type", config.DOCUMENT_TYPES)
    uploaded_files = st.file_uploader(
        "Upload PDF or DOCX", type=["pdf", "docx"],
        accept_multiple_files=True,
    )

    if st.button("➕ Add to Knowledge Base", use_container_width=True):
        if not uploaded_files:
            st.warning("Please select at least one file.")
        else:
            progress = st.progress(0)
            for i, uf in enumerate(uploaded_files):
                with st.spinner(f"Processing {uf.name}…"):
                    try:
                        saved_path = service.save_uploaded_file(uf)
                        result = service.ingest_document(saved_path, doc_type)
                        st.success(f"✅ {uf.name} → {result['num_chunks']} chunks added")
                    except Exception as e:
                        st.error(f"❌ {uf.name}: {e}")
                progress.progress((i + 1) / len(uploaded_files))

    st.divider()
    st.markdown("### Knowledge Base Actions")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔄 Rebuild", use_container_width=True):
            with st.spinner("Rebuilding..."):
                st.cache_resource.clear()
                res = service.rebuild_index()
                st.success(f"{res['total_files']} files, {res['total_chunks']} chunks")
    with col2:
        if st.button("🗑️ Clear All", use_container_width=True):
            if st.session_state.get("confirm_clear"):
                service.clear_knowledge_base()
                st.success("Cleared.")
                st.session_state.confirm_clear = False
            else:
                st.session_state.confirm_clear = True
                st.warning("Click again to confirm.")

    st.divider()
    stats = service.get_stats()
    st.markdown("### 📊 System Status")
    st.metric("Documents", stats["total_documents"])
    st.metric("Chunks", stats["total_chunks"])
    st.markdown(f"**AI Model:** `{stats['ollama_model']}`")

# ══════════════════════════════════════════════════════════════════════════════
# HEADER
# ══════════════════════════════════════════════════════════════════════════════

st.markdown("""
<div class="main-header">
    <h1>📊 AI Marketing Proposal Generator</h1>
    <p>RAG-powered marketing documents from your agency's knowledge base</p>
</div>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# TABS
# ══════════════════════════════════════════════════════════════════════════════

tab1, tab2, tab3, tab4 = st.tabs([
    "📝 Document Generator", "💬 Chat With Documents",
    "📚 Knowledge Base", "🔍 Debug Panel"
])

# ─────────────────────────────────────────────────────────────────────────────
# TAB 1 — DOCUMENT GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

with tab1:
    st.markdown("### Generate a Marketing Document")

    # ── Output type selector ──────────────────────────────────────────────────
    output_type = st.selectbox(
        "📄 What do you want to generate?",
        [
            "📝 Marketing Proposal",
            "📊 Marketing Plan",
            "💰 Pricing Document",
            "📈 Case Study",
            "📋 Service Description",
        ],
        help="Select the type of document to generate"
    )

    st.divider()

    # ── Common fields ─────────────────────────────────────────────────────────
    st.markdown("#### 👤 Client Information")
    col_a, col_b, col_c = st.columns(3)
    with col_a:
        client_name = st.text_input("🏢 Client Name *", placeholder="e.g. Al-Nour Real Estate")
    with col_b:
        industry = st.text_input("🏭 Industry *", placeholder="e.g. Real Estate, Healthcare")
    with col_c:
        language = st.selectbox("🌐 Output Language", ["English", "العربية (Arabic)"])

    col_d, col_e, col_f = st.columns(3)
    with col_d:
        budget = st.text_input("💰 Budget Range", placeholder="e.g. EGP 50,000/month")
    with col_e:
        timeline = st.selectbox("📅 Timeline", ["3 Months", "6 Months", "12 Months", "Custom"])
    with col_f:
        target_audience = st.text_input("👥 Target Audience", placeholder="e.g. Men 25-45, Cairo, Mid-income")

    st.markdown("#### 🎯 Goals & Services")
    col_g, col_h = st.columns(2)
    with col_g:
        goals = st.text_area("🎯 Business Goals", placeholder="e.g. Increase brand awareness, generate 200 leads/month", height=100)
    with col_h:
        services = st.text_area("🛠️ Requested Services", placeholder="e.g. SEO, Google Ads, Social Media, Content Marketing", height=100)

    # ── Dynamic extra fields per output type ──────────────────────────────────
    extra_fields = {}

    if output_type == "📊 Marketing Plan":
        st.markdown("#### 📊 Marketing Plan Details")
        col_i, col_j = st.columns(2)
        with col_i:
            extra_fields["primary_kpi"] = st.selectbox(
                "🎯 Primary KPI",
                ["Lead Generation", "Brand Awareness", "Sales Revenue", "App Installs", "Website Traffic", "Social Media Growth"]
            )
        with col_j:
            extra_fields["current_channels"] = st.text_input(
                "📡 Current Marketing Channels",
                placeholder="e.g. Facebook only, No digital presence"
            )

    elif output_type == "💰 Pricing Document":
        st.markdown("#### 💰 Pricing Details")
        col_i, col_j = st.columns(2)
        with col_i:
            extra_fields["pricing_tier"] = st.selectbox(
                "📦 Pricing Tier",
                ["Basic / Starter", "Standard / Growth", "Premium / Enterprise", "All Tiers"]
            )
        with col_j:
            extra_fields["payment_terms"] = st.selectbox(
                "💳 Payment Terms",
                ["Monthly", "Quarterly", "Annual", "Project-based"]
            )

    elif output_type == "📈 Case Study":
        st.markdown("#### 📈 Case Study Details")
        col_i, col_j = st.columns(2)
        with col_i:
            extra_fields["results_achieved"] = st.text_area(
                "📊 Results Achieved",
                placeholder="e.g. 45% increase in leads, Cost per lead reduced from EGP 800 to EGP 320",
                height=80
            )
        with col_j:
            extra_fields["campaign_duration"] = st.text_input(
                "⏱️ Campaign Duration",
                placeholder="e.g. 6 months (Jan-Jun 2024)"
            )

    elif output_type == "📋 Service Description":
        st.markdown("#### 📋 Service Details")
        extra_fields["service_focus"] = st.multiselect(
            "🛠️ Services to Describe",
            ["SEO", "Google Ads", "Meta Ads", "TikTok Ads", "Social Media Management",
             "Content Marketing", "Email Marketing", "Influencer Marketing", "CRM Setup",
             "Website Development", "Video Production", "Photography"],
            default=["SEO", "Google Ads", "Social Media Management"]
        )

    st.divider()

    # ── Generate button ───────────────────────────────────────────────────────
    generate_btn = st.button(
        f"🚀 Generate {output_type.split(' ', 1)[1]}",
        use_container_width=True,
        type="primary"
    )

    if generate_btn:
        if not client_name or not industry:
            st.error("Please provide at least a Client Name and Industry.")
        elif stats["total_chunks"] == 0:
            st.warning("⚠️ Knowledge base is empty. Upload documents first.")
        else:
            lang_instruction = "Respond entirely in Arabic." if "Arabic" in language else "Respond in English."

            # Build prompt based on output type
            if output_type == "📝 Marketing Proposal":
                prompt_extra = f"""
Generate a complete MARKETING PROPOSAL with these sections:
1. Executive Summary | 2. Client Background & Goals | 3. Recommended Services
4. Scope of Work | 5. Deliverables | 6. Timeline ({timeline})
7. Pricing Recommendation | 8. KPIs | 9. Next Steps
{lang_instruction}"""

            elif output_type == "📊 Marketing Plan":
                prompt_extra = f"""
Generate a complete {timeline} MARKETING PLAN with these sections:
1. Executive Summary | 2. Market Analysis | 3. Target Audience Analysis
4. Channel Strategy | 5. Monthly Activity Calendar | 6. Budget Allocation
7. KPIs & Measurement (Primary KPI: {extra_fields.get('primary_kpi', '')})
8. Implementation Timeline
Current channels: {extra_fields.get('current_channels', 'Not specified')}
{lang_instruction}"""

            elif output_type == "💰 Pricing Document":
                prompt_extra = f"""
Generate a detailed PRICING DOCUMENT including:
1. Service Packages ({extra_fields.get('pricing_tier', 'All Tiers')})
2. What's Included in Each Package | 3. Pricing Table with Clear Numbers
4. Add-on Services & Pricing | 5. Payment Terms ({extra_fields.get('payment_terms', 'Monthly')})
6. ROI Expectations per Package | 7. Terms & Conditions
{lang_instruction}"""

            elif output_type == "📈 Case Study":
                prompt_extra = f"""
Generate a compelling CASE STUDY with these sections:
1. Client Overview | 2. The Challenge | 3. Our Strategy & Approach
4. Implementation Details | 5. Results & Impact
Results achieved: {extra_fields.get('results_achieved', 'Not specified')}
Campaign duration: {extra_fields.get('campaign_duration', 'Not specified')}
6. Key Learnings | 7. Client Testimonial (if available)
{lang_instruction}"""

            elif output_type == "📋 Service Description":
                services_list = ", ".join(extra_fields.get('service_focus', [services]))
                prompt_extra = f"""
Generate a professional SERVICE DESCRIPTION document for: {services_list}
Include for each service:
1. What It Is | 2. What We Do (Deliverables) | 3. Who It's For
4. Expected Results & Timeline | 5. Our Process | 6. Pricing Range
{lang_instruction}"""

            with st.spinner(f"🔍 Searching knowledge base and generating {output_type.split(' ', 1)[1]}…"):
                try:
                    query = f"{industry} {output_type} {goals} {services} {target_audience}"
                    retrieved = service._retriever.retrieve(query)
                    context = service._retriever.format_context(retrieved)

                    full_prompt = f"""
## CLIENT BRIEF
- Client: {client_name}
- Industry: {industry}
- Budget: {budget}
- Timeline: {timeline}
- Target Audience: {target_audience}
- Business Goals: {goals}
- Requested Services: {services}

## CONTEXT FROM KNOWLEDGE BASE
{context}

## TASK
{prompt_extra}
"""
                    result = service._llm.generate_response(full_prompt)

                    st.session_state.last_retrieved_chunks = retrieved
                    st.session_state.last_context_sent = f"{output_type} for: {client_name} / {industry}"

                    st.success(f"✅ {output_type.split(' ', 1)[1]} generated using {len(retrieved)} knowledge base chunks.")
                    st.divider()

                    # Download buttons
                    def md_to_docx(text, title):
                        doc = DocxDocument()
                        doc.add_heading(title, 0)
                        for ln in text.split("\n"):
                            ln = ln.strip()
                            if not ln: doc.add_paragraph("")
                            elif ln.startswith("### "): doc.add_heading(ln[4:], level=2)
                            elif ln.startswith("## "): doc.add_heading(ln[3:], level=1)
                            elif ln.startswith("# "): doc.add_heading(ln[2:], level=0)
                            elif ln.startswith("- ") or ln.startswith("• "): doc.add_paragraph(ln[2:], style="List Bullet")
                            else: doc.add_paragraph(ln)
                        buf = io.BytesIO()
                        doc.save(buf)
                        buf.seek(0)
                        return buf.getvalue()

                    doc_title = f"{output_type.split(' ', 1)[1]} - {client_name}"
                    col_dl1, col_dl2 = st.columns(2)
                    with col_dl1:
                        st.download_button(
                            "⬇️ Download Word",
                            data=md_to_docx(result, doc_title),
                            file_name=f"{client_name.replace(' ', '_')}_{output_type.split(' ', 1)[1].replace(' ', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    with col_dl2:
                        st.download_button(
                            "⬇️ Download Markdown",
                            data=result,
                            file_name=f"{client_name.replace(' ', '_')}_{output_type.split(' ', 1)[1].replace(' ', '_')}.md",
                            mime="text/markdown",
                            use_container_width=True
                        )

                    st.markdown(result)

                except Exception as e:
                    st.error(f"❌ Error: {e}")
                    st.code(traceback.format_exc())

# ─────────────────────────────────────────────────────────────────────────────
# TAB 2 — CHAT WITH DOCUMENTS
# ─────────────────────────────────────────────────────────────────────────────

with tab2:
    st.markdown("### 💬 Chat With Your Document Library")
    st.caption("Ask questions about your previous proposals, case studies, and pricing documents.")

    example_questions = [
        "Show me similar proposals for real estate clients.",
        "What pricing models have we used before?",
        "What services are included in our SEO packages?",
        "What KPIs were proposed for healthcare clients?",
        "ما هي الخدمات التسويقية المقدمة للمطاعم؟",
    ]

    st.markdown("**💡 Example questions:**")
    cols = st.columns(len(example_questions))
    for col, q in zip(cols, example_questions):
        if col.button(q, use_container_width=True, key=f"ex_{q[:20]}"):
            st.session_state.chat_history.append({"role": "user", "content": q})
            with st.spinner("Searching documents…"):
                answer, retrieved = service.chat(q)
                st.session_state.chat_history.append({"role": "assistant", "content": answer})
                st.session_state.last_retrieved_chunks = retrieved
                st.session_state.last_context_sent = f"Chat: {q}"

    for msg in st.session_state.chat_history:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    if user_input := st.chat_input("Ask anything about your documents… (Arabic or English)"):
        st.session_state.chat_history.append({"role": "user", "content": user_input})
        with st.chat_message("user"):
            st.markdown(user_input)

        if stats["total_chunks"] == 0:
            response = "⚠️ Knowledge base is empty. Please upload documents first."
            retrieved = []
        else:
            with st.spinner("Searching knowledge base…"):
                response, retrieved = service.chat(user_input)
                st.session_state.last_retrieved_chunks = retrieved
                st.session_state.last_context_sent = f"Chat: {user_input}"

        st.session_state.chat_history.append({"role": "assistant", "content": response})
        with st.chat_message("assistant"):
            st.markdown(response)

    if st.session_state.chat_history:
        if st.button("🗑️ Clear Chat History"):
            st.session_state.chat_history = []
            st.rerun()

# ─────────────────────────────────────────────────────────────────────────────
# TAB 3 — KNOWLEDGE BASE
# ─────────────────────────────────────────────────────────────────────────────

with tab3:
    st.markdown("### 📚 Knowledge Base Overview")
    if st.button("🔄 Refresh Stats"):
        st.rerun()

    stats = service.get_stats()
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("📄 Documents", stats["total_documents"])
    c2.metric("🧩 Chunks", stats["total_chunks"])
    c3.metric("🤖 AI Model", stats["ollama_model"].split("/")[-1] if "/" in stats["ollama_model"] else stats["ollama_model"])
    c4.metric("💾 FAISS", "Active" if stats["total_chunks"] > 0 else "Empty")

    st.divider()

    if stats["total_chunks"] == 0:
        st.info("No documents in the knowledge base yet. Upload files from the sidebar.")
    else:
        col_left, col_right = st.columns(2)
        with col_left:
            st.markdown("#### 📁 Documents")
            for doc_name, chunk_count in stats["documents"].items():
                st.markdown(f"- **{doc_name}** — {chunk_count} chunks")
        with col_right:
            st.markdown("#### 🏷️ Document Types")
            for dtype, count in stats["document_types"].items():
                st.markdown(f"- **{dtype}** — {count} chunks")

# ─────────────────────────────────────────────────────────────────────────────
# TAB 4 — DEBUG PANEL
# ─────────────────────────────────────────────────────────────────────────────

with tab4:
    st.markdown("### 🔍 Developer Debug Panel")
    retrieved = st.session_state.last_retrieved_chunks

    if not retrieved:
        st.info("No retrieval data yet. Generate a document or ask a chat question first.")
    else:
        st.markdown(f"**Last query:** `{st.session_state.last_context_sent}`")
        st.markdown(f"**Chunks retrieved:** {len(retrieved)}")
        st.divider()

        for i, chunk in enumerate(retrieved, start=1):
            score = chunk["score"]
            meta = chunk["metadata"]
            score_class = "score-high" if score >= 0.7 else "score-mid" if score >= 0.4 else "score-low"

            with st.expander(
                f"Chunk {i} | {meta.get('source_file', '?')} | Score: {score:.4f}",
                expanded=(i == 1),
            ):
                col_meta, col_score = st.columns([3, 1])
                with col_meta:
                    st.markdown(f"**Source:** `{meta.get('source_file', 'unknown')}`")
                    st.markdown(f"**Section:** `{meta.get('section_name', 'unknown')}`")
                    st.markdown(f"**Type:** `{meta.get('document_type', 'unknown')}`")
                with col_score:
                    st.markdown(f"<p class='{score_class}'>Score: {score:.4f}</p>", unsafe_allow_html=True)
                st.text(chunk["text"])
