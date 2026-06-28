"""
Document Loader
===============
Loads PDF and DOCX files and returns raw extracted text.
Keeps I/O concerns separate from chunking / cleaning concerns.
"""

from __future__ import annotations

import logging
from pathlib import Path

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Extracts plain text from PDF and DOCX files."""

    SUPPORTED_EXTENSIONS: set[str] = {".pdf", ".docx"}

    # ── Public API ─────────────────────────────────────────────────────────────

    def load(self, file_path: str | Path) -> str:
        """
        Load a document and return its full text content.

        Args:
            file_path: Path to a PDF or DOCX file.

        Returns:
            Extracted text as a single string.

        Raises:
            ValueError: If the file extension is not supported.
            FileNotFoundError: If the file does not exist.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type '{suffix}'. "
                f"Supported: {self.SUPPORTED_EXTENSIONS}"
            )

        logger.info("Loading document: %s", path.name)

        if suffix == ".pdf":
            return self._load_pdf(path)
        return self._load_docx(path)

    # ── Private helpers ────────────────────────────────────────────────────────

    def _load_pdf(self, path: Path) -> str:
        """Extract text from a PDF using pdfplumber."""
        pages: list[str] = []
        try:
            with pdfplumber.open(path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text()
                    if text:
                        pages.append(text)
                    else:
                        logger.debug("Page %d of %s yielded no text.", page_num, path.name)
        except Exception as exc:
            logger.error("Failed to read PDF '%s': %s", path.name, exc)
            raise

        full_text = "\n\n".join(pages)
        logger.info("Extracted %d characters from PDF '%s'.", len(full_text), path.name)
        return full_text

    def _load_docx(self, path: Path) -> str:
        """Extract text from a DOCX file paragraph by paragraph."""
        try:
            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also pull text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
        except Exception as exc:
            logger.error("Failed to read DOCX '%s': %s", path.name, exc)
            raise

        full_text = "\n\n".join(paragraphs)
        logger.info("Extracted %d characters from DOCX '%s'.", len(full_text), path.name)
        return full_text
